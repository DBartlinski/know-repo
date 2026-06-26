"""
Add metadata headers to DOCX and PDF files from query.xlsx metadata.
"""

import os
import pandas as pd
from pathlib import Path
from datetime import datetime
from typing import Dict, Optional, Tuple
import logging
from io import BytesIO
import tempfile

# Import document processing libraries
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement

# Import PDF processing libraries
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
from reportlab.lib.units import inch
from reportlab.lib.colors import HexColor
from pypdf import PdfReader, PdfWriter

# Setup logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

# Configuration
QUERY_XLSX_PATH = r'C:\Users\vhacobartld\VS Code\RFIs\query.xlsx'
RFI_DOCS_PATH = r'C:\Users\vhacobartld\VS Code\RFIs\RFI-Docs'
METADATA_FONT_SIZE = 9
METADATA_FONT_COLOR = RGBColor(100, 100, 100)  # Dark gray


class MetadataProcessor:
    def __init__(self, xlsx_path: str, docs_folder: str):
        """Initialize the metadata processor."""
        self.xlsx_path = xlsx_path
        self.docs_folder = docs_folder
        self.metadata_map: Dict[str, Dict] = {}
        self.processed_files = 0
        self.skipped_files = 0
        self.failed_files = []
        
    def load_metadata(self) -> bool:
        """Load metadata from query.xlsx file."""
        try:
            logger.info(f"Loading metadata from {self.xlsx_path}")
            df = pd.read_excel(self.xlsx_path, sheet_name='query')
            
            # Create a map of filename -> metadata
            for idx, row in df.iterrows():
                filename = str(row['Name']).strip()
                self.metadata_map[filename] = {
                    'name': filename,
                    'title': str(row['Title']).strip() if pd.notna(row['Title']) else '',
                    'topic': str(row['Topic']).strip() if pd.notna(row['Topic']) else '',
                    'document_type': str(row['Document Type']).strip() if pd.notna(row['Document Type']) else '',
                    'created': row['Created'] if pd.notna(row['Created']) else None,
                    'created_by': str(row['Created By']).strip() if pd.notna(row['Created By']) else '',
                    'modified': row['Modified'] if pd.notna(row['Modified']) else None,
                    'modified_by': str(row['Modified By']).strip() if pd.notna(row['Modified By']) else '',
                }
            
            logger.info(f"Loaded metadata for {len(self.metadata_map)} documents")
            return True
        except Exception as e:
            logger.error(f"Error loading metadata: {e}")
            return False
    
    def format_metadata_header(self, metadata: Dict) -> list:
        """Format metadata into header lines."""
        lines = []
        
        # Line 1: Document Name
        if metadata['name']:
            lines.append(f"Document: {metadata['name']}")
        
        # Line 2: Title
        if metadata['title']:
            lines.append(f"Title: {metadata['title']}")
        
        # Line 3: Topic & Type
        topic_type = []
        if metadata['topic']:
            topic_type.append(f"Topic: {metadata['topic']}")
        if metadata['document_type']:
            topic_type.append(f"Type: {metadata['document_type']}")
        if topic_type:
            lines.append(" | ".join(topic_type))
        
        # Line 4: Created info
        created_info = []
        if metadata['created']:
            created_date = metadata['created'].strftime('%Y-%m-%d')
            created_info.append(f"Created: {created_date}")
        if metadata['created_by']:
            created_info.append(f"By: {metadata['created_by']}")
        if created_info:
            lines.append(" | ".join(created_info))
        
        # Line 5: Modified info
        if metadata['modified'] or metadata['modified_by']:
            modified_info = []
            if metadata['modified']:
                modified_date = metadata['modified'].strftime('%Y-%m-%d')
                modified_info.append(f"Modified: {modified_date}")
            if metadata['modified_by']:
                modified_info.append(f"By: {metadata['modified_by']}")
            if modified_info:
                lines.append(" | ".join(modified_info))
        
        return lines
    
    def add_header_to_docx(self, file_path: str, metadata: Dict) -> bool:
        """Add metadata header to DOCX file."""
        try:
            doc = Document(file_path)
            
            # Get header lines
            header_lines = self.format_metadata_header(metadata)
            
            # Insert paragraphs at the beginning
            body = doc._body._element
            
            # Insert each line as a paragraph at the beginning (in reverse order so they end up in correct order)
            for i, line in enumerate(reversed(header_lines)):
                # Create a new paragraph element
                new_p = OxmlElement('w:p')
                # Insert at position 0
                body.insert(0, new_p)
            
            # Now format the header paragraphs
            for i, line in enumerate(header_lines):
                para = doc.paragraphs[i]
                para.text = line
                para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                
                # Format the text
                for run in para.runs:
                    run.font.size = Pt(METADATA_FONT_SIZE)
                    run.font.color.rgb = METADATA_FONT_COLOR
            
            # Add separator after header
            sep_para = doc.paragraphs[len(header_lines)]
            sep_para.text = "─" * 80
            sep_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            sep_run = sep_para.runs[0]
            sep_run.font.size = Pt(8)
            sep_run.font.color.rgb = RGBColor(200, 200, 200)
            
            # Add spacing after separator
            spacing_para = doc.paragraphs[len(header_lines) + 1]
            spacing_para.text = ""
            
            # Save the document
            doc.save(file_path)
            logger.info(f"✓ Added header to: {Path(file_path).name}")
            return True
        
        except Exception as e:
            logger.error(f"✗ Error processing DOCX {Path(file_path).name}: {e}")
            self.failed_files.append((Path(file_path).name, str(e)))
            return False
    
    def add_header_to_pdf(self, file_path: str, metadata: Dict) -> bool:
        """Add metadata header to PDF file by prepending a cover page."""
        try:
            # Create a temporary PDF with metadata cover page
            temp_buffer = BytesIO()
            
            # Create the metadata cover page using reportlab
            doc = SimpleDocTemplate(
                temp_buffer,
                pagesize=letter,
                rightMargin=0.75*inch,
                leftMargin=0.75*inch,
                topMargin=1*inch,
                bottomMargin=1*inch
            )
            
            # Create styles
            styles = getSampleStyleSheet()
            title_style = ParagraphStyle(
                'MetadataTitle',
                parent=styles['Heading1'],
                fontSize=16,
                textColor=HexColor('333333'),
                spaceAfter=12,
                alignment=0  # Left align
            )
            
            normal_style = ParagraphStyle(
                'MetadataBody',
                parent=styles['Normal'],
                fontSize=10,
                textColor=HexColor('666666'),
                spaceAfter=6,
                alignment=0  # Left align
            )
            
            label_style = ParagraphStyle(
                'MetadataLabel',
                parent=styles['Normal'],
                fontSize=9,
                textColor=HexColor('888888'),
                spaceAfter=4,
                alignment=0  # Left align
            )
            
            # Build the story (content) for the cover page
            story = []
            story.append(Paragraph("DOCUMENT METADATA", title_style))
            story.append(Spacer(1, 0.2*inch))
            
            # Add metadata fields
            if metadata['name']:
                story.append(Paragraph(f"<b>Document Name:</b>", label_style))
                story.append(Paragraph(metadata['name'], normal_style))
                story.append(Spacer(1, 0.1*inch))
            
            if metadata['title']:
                story.append(Paragraph(f"<b>Title:</b>", label_style))
                story.append(Paragraph(metadata['title'], normal_style))
                story.append(Spacer(1, 0.1*inch))
            
            if metadata['topic']:
                story.append(Paragraph(f"<b>Topic:</b>", label_style))
                story.append(Paragraph(metadata['topic'], normal_style))
                story.append(Spacer(1, 0.1*inch))
            
            if metadata['document_type']:
                story.append(Paragraph(f"<b>Document Type:</b>", label_style))
                story.append(Paragraph(metadata['document_type'], normal_style))
                story.append(Spacer(1, 0.1*inch))
            
            if metadata['created']:
                created_date = metadata['created'].strftime('%Y-%m-%d')
                story.append(Paragraph(f"<b>Created:</b>", label_style))
                created_info = f"{created_date}"
                if metadata['created_by']:
                    created_info += f" by {metadata['created_by']}"
                story.append(Paragraph(created_info, normal_style))
                story.append(Spacer(1, 0.1*inch))
            
            if metadata['modified']:
                modified_date = metadata['modified'].strftime('%Y-%m-%d')
                story.append(Paragraph(f"<b>Modified:</b>", label_style))
                modified_info = f"{modified_date}"
                if metadata['modified_by']:
                    modified_info += f" by {metadata['modified_by']}"
                story.append(Paragraph(modified_info, normal_style))
                story.append(Spacer(1, 0.1*inch))
            
            # Build the PDF
            doc.build(story)
            
            # Get the cover page PDF
            temp_buffer.seek(0)
            cover_pdf = PdfReader(temp_buffer)
            
            # Read the original PDF
            original_pdf = PdfReader(file_path)
            
            # Create output PDF
            output_pdf = PdfWriter()
            
            # Add the cover page first
            output_pdf.add_page(cover_pdf.pages[0])
            
            # Add all pages from the original PDF
            for page in original_pdf.pages:
                output_pdf.add_page(page)
            
            # Write the output PDF back to the original file
            with open(file_path, 'wb') as output_file:
                output_pdf.write(output_file)
            
            logger.info(f"✓ Added metadata cover page to: {Path(file_path).name}")
            return True
        
        except Exception as e:
            logger.error(f"✗ Error processing PDF {Path(file_path).name}: {e}")
            self.failed_files.append((Path(file_path).name, str(e)))
            return False
    
    def process_documents(self) -> Tuple[int, int, list]:
        """Process all documents in RFI-Docs folder."""
        if not os.path.exists(self.docs_folder):
            logger.error(f"RFI-Docs folder not found: {self.docs_folder}")
            return 0, 0, []
        
        logger.info(f"Processing documents in {self.docs_folder}")
        
        for filename in os.listdir(self.docs_folder):
            file_path = os.path.join(self.docs_folder, filename)
            
            # Skip directories
            if os.path.isdir(file_path):
                continue
            
            # Check if we have metadata for this file
            if filename not in self.metadata_map:
                logger.debug(f"No metadata found for: {filename}")
                continue
            
            metadata = self.metadata_map[filename]
            
            # Process based on file extension
            if filename.lower().endswith('.docx'):
                if self.add_header_to_docx(file_path, metadata):
                    self.processed_files += 1
            elif filename.lower().endswith('.pdf'):
                if self.add_header_to_pdf(file_path, metadata):
                    self.processed_files += 1
                else:
                    self.skipped_files += 1
        
        return self.processed_files, self.skipped_files, self.failed_files
    
    def print_summary(self):
        """Print processing summary."""
        print("\n" + "=" * 60)
        print("METADATA HEADER INJECTION SUMMARY")
        print("=" * 60)
        print(f"Successfully processed: {self.processed_files} files")
        print(f"  (DOCX files with text headers + PDF files with cover pages)")
        print(f"Skipped: {self.skipped_files} files")
        
        if self.failed_files:
            print(f"\nFailed files ({len(self.failed_files)}):")
            for filename, error in self.failed_files:
                print(f"  - {filename}")
                print(f"    Error: {error}")
        
        print("=" * 60)


def main():
    """Main entry point."""
    processor = MetadataProcessor(QUERY_XLSX_PATH, RFI_DOCS_PATH)
    
    # Load metadata
    if not processor.load_metadata():
        logger.error("Failed to load metadata. Exiting.")
        return False
    
    # Process documents
    processed, skipped, failed = processor.process_documents()
    
    # Print summary
    processor.print_summary()
    
    return len(failed) == 0


if __name__ == '__main__':
    import sys
    success = main()
    sys.exit(0 if success else 1)
