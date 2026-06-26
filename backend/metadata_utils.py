"""
Reusable metadata utilities for adding headers to documents and managing metadata.
"""

from pathlib import Path
from datetime import datetime
from typing import Dict
from io import BytesIO
import logging

# Import document processing libraries
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement

# Import PDF processing libraries
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.units import inch
from reportlab.lib.colors import HexColor
from pypdf import PdfReader, PdfWriter

logger = logging.getLogger(__name__)

# Configuration
METADATA_FONT_SIZE = 9
METADATA_FONT_COLOR = RGBColor(100, 100, 100)  # Dark gray


def format_metadata_header(metadata: Dict) -> list:
    """Format metadata into header lines."""
    lines = []
    
    # Line 1: Document Name
    if metadata.get('name'):
        lines.append(f"Document: {metadata['name']}")
    
    # Line 2: Title
    if metadata.get('title'):
        lines.append(f"Title: {metadata['title']}")
    
    # Line 3: Topic & Type
    topic_type = []
    if metadata.get('topic'):
        topic_type.append(f"Topic: {metadata['topic']}")
    if metadata.get('doc_type') or metadata.get('document_type'):
        doc_type = metadata.get('doc_type') or metadata.get('document_type')
        topic_type.append(f"Type: {doc_type}")
    if topic_type:
        lines.append(" | ".join(topic_type))
    
    # Line 4: Created info
    created_info = []
    if metadata.get('created'):
        if isinstance(metadata['created'], str):
            created_date = metadata['created']
        else:
            created_date = metadata['created'].strftime('%Y-%m-%d') if hasattr(metadata['created'], 'strftime') else str(metadata['created'])
        created_info.append(f"Created: {created_date}")
    if metadata.get('created_by'):
        created_info.append(f"By: {metadata['created_by']}")
    if created_info:
        lines.append(" | ".join(created_info))
    
    # Line 5: Modified info
    if metadata.get('modified') or metadata.get('modified_by'):
        modified_info = []
        if metadata.get('modified'):
            if isinstance(metadata['modified'], str):
                modified_date = metadata['modified']
            else:
                modified_date = metadata['modified'].strftime('%Y-%m-%d') if hasattr(metadata['modified'], 'strftime') else str(metadata['modified'])
            modified_info.append(f"Modified: {modified_date}")
        if metadata.get('modified_by'):
            modified_info.append(f"By: {metadata['modified_by']}")
        if modified_info:
            lines.append(" | ".join(modified_info))
    
    return lines


def add_header_to_docx(file_path: str, metadata: Dict) -> bool:
    """Add metadata header to DOCX file."""
    try:
        doc = Document(file_path)
        
        # Get header lines
        header_lines = format_metadata_header(metadata)
        
        # Insert paragraphs at the beginning
        body = doc._body._element
        
        # Insert each line as a paragraph at the beginning (in reverse order so they end up in correct order)
        for i, line in enumerate(reversed(header_lines)):
            # Create a new paragraph element
            new_p = OxmlElement('w:p')
            # Insert at position 0
            body.insert(0, new_p)
        
        # Now format the header paragraphs
        for i, line in enumerate(header_lines):
            para = doc.paragraphs[i]
            para.text = line
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            
            # Format the text
            for run in para.runs:
                run.font.size = Pt(METADATA_FONT_SIZE)
                run.font.color.rgb = METADATA_FONT_COLOR
        
        # Add separator after header
        sep_para = doc.paragraphs[len(header_lines)]
        sep_para.text = "─" * 80
        sep_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        sep_run = sep_para.runs[0]
        sep_run.font.size = Pt(8)
        sep_run.font.color.rgb = RGBColor(200, 200, 200)
        
        # Add spacing after separator
        spacing_para = doc.paragraphs[len(header_lines) + 1]
        spacing_para.text = ""
        
        # Save the document
        doc.save(file_path)
        logger.info(f"✓ Added header to: {Path(file_path).name}")
        return True
    
    except Exception as e:
        logger.error(f"✗ Error processing DOCX {Path(file_path).name}: {e}")
        return False


def add_header_to_pdf(file_path: str, metadata: Dict) -> bool:
    """Add metadata header to PDF file by prepending a cover page."""
    try:
        # Create a temporary PDF with metadata cover page
        temp_buffer = BytesIO()
        
        # Create the metadata cover page using reportlab
        doc = SimpleDocTemplate(
            temp_buffer,
            pagesize=letter,
            rightMargin=0.75*inch,
            leftMargin=0.75*inch,
            topMargin=1*inch,
            bottomMargin=1*inch
        )
        
        # Create styles
        styles = getSampleStyleSheet()
        title_style = ParagraphStyle(
            'MetadataTitle',
            parent=styles['Heading1'],
            fontSize=16,
            textColor=HexColor('333333'),
            spaceAfter=12,
            alignment=0  # Left align
        )
        
        normal_style = ParagraphStyle(
            'MetadataBody',
            parent=styles['Normal'],
            fontSize=10,
            textColor=HexColor('666666'),
            spaceAfter=6,
            alignment=0  # Left align
        )
        
        label_style = ParagraphStyle(
            'MetadataLabel',
            parent=styles['Normal'],
            fontSize=9,
            textColor=HexColor('888888'),
            spaceAfter=4,
            alignment=0  # Left align
        )
        
        # Build the story (content) for the cover page
        story = []
        story.append(Paragraph("DOCUMENT METADATA", title_style))
        story.append(Spacer(1, 0.2*inch))
        
        # Add metadata fields
        if metadata.get('name'):
            story.append(Paragraph(f"<b>Document Name:</b>", label_style))
            story.append(Paragraph(metadata['name'], normal_style))
            story.append(Spacer(1, 0.1*inch))
        
        if metadata.get('title'):
            story.append(Paragraph(f"<b>Title:</b>", label_style))
            story.append(Paragraph(metadata['title'], normal_style))
            story.append(Spacer(1, 0.1*inch))
        
        if metadata.get('topic'):
            story.append(Paragraph(f"<b>Topic:</b>", label_style))
            story.append(Paragraph(metadata['topic'], normal_style))
            story.append(Spacer(1, 0.1*inch))
        
        doc_type = metadata.get('doc_type') or metadata.get('document_type')
        if doc_type:
            story.append(Paragraph(f"<b>Document Type:</b>", label_style))
            story.append(Paragraph(doc_type, normal_style))
            story.append(Spacer(1, 0.1*inch))
        
        if metadata.get('created'):
            if isinstance(metadata['created'], str):
                created_date = metadata['created']
            else:
                created_date = metadata['created'].strftime('%Y-%m-%d') if hasattr(metadata['created'], 'strftime') else str(metadata['created'])
            story.append(Paragraph(f"<b>Created:</b>", label_style))
            created_info = f"{created_date}"
            if metadata.get('created_by'):
                created_info += f" by {metadata['created_by']}"
            story.append(Paragraph(created_info, normal_style))
            story.append(Spacer(1, 0.1*inch))
        
        if metadata.get('modified'):
            if isinstance(metadata['modified'], str):
                modified_date = metadata['modified']
            else:
                modified_date = metadata['modified'].strftime('%Y-%m-%d') if hasattr(metadata['modified'], 'strftime') else str(metadata['modified'])
            story.append(Paragraph(f"<b>Modified:</b>", label_style))
            modified_info = f"{modified_date}"
            if metadata.get('modified_by'):
                modified_info += f" by {metadata['modified_by']}"
            story.append(Paragraph(modified_info, normal_style))
            story.append(Spacer(1, 0.1*inch))
        
        # Build the PDF
        doc.build(story)
        
        # Get the cover page PDF
        temp_buffer.seek(0)
        cover_pdf = PdfReader(temp_buffer)
        
        # Read the original PDF
        original_pdf = PdfReader(file_path)
        
        # Create output PDF
        output_pdf = PdfWriter()
        
        # Add the cover page first
        output_pdf.add_page(cover_pdf.pages[0])
        
        # Add all pages from the original PDF
        for page in original_pdf.pages:
            output_pdf.add_page(page)
        
        # Write the output PDF back to the original file
        with open(file_path, 'wb') as output_file:
            output_pdf.write(output_file)
        
        logger.info(f"✓ Added metadata cover page to: {Path(file_path).name}")
        return True
    
    except Exception as e:
        logger.error(f"✗ Error processing PDF {Path(file_path).name}: {e}")
        return False
