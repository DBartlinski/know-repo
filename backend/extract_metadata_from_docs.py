"""
Extract metadata (Topic and Document Type) from document headers/content
and populate the database for filtering and search.

This script uses a two-tier approach:
1. PRIMARY: Extract from structured headers (DOCX paragraphs or PDF cover page)
2. FALLBACK: Analyze document content if headers missing/incomplete
   - Search for keywords matching predefined topics
   - Infer document type from content patterns
"""

import os
import re
import sqlite3
import logging
from pathlib import Path
from typing import Dict, Optional, Tuple, Set

from docx import Document
from pypdf import PdfReader

from config import DB_PATH, DOC_FOLDER, PREDEFINED_TOPICS, PREDEFINED_DOCUMENT_TYPES

# Setup logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)


class MetadataExtractor:
    def __init__(self, db_path: str, docs_folder: str):
        """Initialize the metadata extractor."""
        self.db_path = db_path
        self.docs_folder = docs_folder
        self.processed_count = 0
        self.skipped_count = 0
        self.failed_count = 0
        self.failed_files = []
        
        # Build topic keyword map for fallback extraction
        self.topic_keywords = self._build_topic_keywords()
        
    def _build_topic_keywords(self) -> Dict[str, Set[str]]:
        """Build mapping of topics to their associated keywords."""
        return {
            'Psychedelics': {'psychedelic', 'psilocybin', 'mdma', 'lsd', 'magic mushroom', 'ecstasy'},
            'Mental Health': {'ptsd', 'depression', 'anxiety', 'mental health', 'suicide', 'trauma'},
            'Traumatic Brain Injury (TBI)': {'tbi', 'traumatic brain', 'mild traumatic', 'concussion', 'blast injury'},
            'Precision Oncology': {'cancer', 'oncology', 'tumor', 'chemotherapy', 'precision medicine'},
            'Clinical Trials': {'clinical trial', 'randomized controlled trial', 'rct', 'study protocol'},
            'Women\'s Health': {'women veteran', 'women\'s health', 'pregnancy', 'reproductive'},
            'Rural Health': {'rural', 'rural veteran', 'rural healthcare', 'geographic disparity'},
            'MVP': {'million veteran', 'mvp', 'genetic', 'genomic'},
            'Animal Research': {'animal research', 'animal model', 'laboratory animal', 'preclinical'},
            'Congress': {'congress', 'congressional', 'hearing', 'senate', 'representative'},
            'COVID-19': {'covid', 'covid-19', 'pandemic', 'coronavirus'},
            'Cannabis': {'cannabis', 'marijuana', 'hemp', 'thc', 'cbd'},
        }
    
    def _infer_doc_type_from_content(self, text: str) -> Optional[str]:
        """Infer document type from content patterns."""
        text_lower = text.lower()
        
        # Research/Study patterns
        if any(phrase in text_lower for phrase in ['study shows', 'research found', 'research demonstrates', 
                                                      'published in', 'journal of', 'authors:', 'abstract']):
            if any(phrase in text_lower for phrase in ['conclusion', 'methodology', 'results']):
                return 'Publications-Studies'
        
        # News/Article patterns
        if any(phrase in text_lower for phrase in ['news', 'article', 'press', 'reported', 'announced', 
                                                      'said', 'told', 'staff writer', 'published']):
            if 'statement' not in text_lower:
                return 'News Articles'
        
        # Press Release
        if any(phrase in text_lower for phrase in ['press release', 'for immediate release', 'media contact']):
            return 'Press Releases'
        
        # Testimony/Statement patterns
        if any(phrase in text_lower for phrase in ['statement from', 'joint statement', 'oral testimony', 
                                                      'statement before', 'testimony', 'hearing', 'statement on behalf']):
            return 'Testimony and Statements'
        
        # Briefing/Talking Points
        if any(phrase in text_lower for phrase in ['talking points', 'briefing', 'key messages', 'background']):
            return 'Briefing Papers and Congressional Documents'
        
        # Presentation patterns
        if any(phrase in text_lower for phrase in ['presentation', 'slide', 'agenda', 'program overview']):
            return 'Presentations and Other Documents'
        
        # Report patterns
        if any(phrase in text_lower for phrase in ['annual report', 'final report', 'report', 'findings', 
                                                      'recommendations', 'executive summary']):
            return 'Reports'
        
        # RFA/RFI patterns
        if any(phrase in text_lower for phrase in ['request for application', 'rfa', 'request for information', 'rfi']):
            return 'RFAs and Research'
        
        # Default to Other-Misc
        return 'Other-Misc Documents'
    
    def get_connection(self):
        """Get database connection with row factory."""
        conn = sqlite3.connect(self.db_path)
        conn.row_factory = sqlite3.Row
        return conn
    
    def extract_metadata_from_docx(self, file_path: str) -> Dict[str, Optional[str]]:
        """Extract Topic and Document Type from DOCX metadata header (primary) or content (fallback)."""
        try:
            doc = Document(file_path)
            metadata = {'topic': None, 'document_type': None}
            
            # PRIMARY: Try to extract from structured header (first 8 paragraphs)
            for i, para in enumerate(doc.paragraphs[:8]):
                text = para.text.strip()
                
                # Look for combined line "Topic: X | Type: Y"
                if 'Topic:' in text and 'Type:' in text:
                    match = re.search(r'Topic:\s*([^|]+)\s*\|\s*Type:\s*(.+)', text)
                    if match:
                        metadata['topic'] = match.group(1).strip()
                        metadata['document_type'] = match.group(2).strip()
                        if metadata['topic'] and metadata['document_type']:
                            return metadata
                
                # Look for Topic line
                if text.startswith('Topic:'):
                    topic_part = text.split('|')[0].replace('Topic:', '').strip()
                    if topic_part:
                        metadata['topic'] = topic_part
                
                # Look for Type line
                if text.startswith('Type:'):
                    doc_type = text.replace('Type:', '').strip()
                    if doc_type:
                        metadata['document_type'] = doc_type
            
            # If we got both from header, return
            if metadata['topic'] and metadata['document_type']:
                return metadata
            
            # FALLBACK: Analyze document content if header incomplete
            full_text = '\n'.join([p.text for p in doc.paragraphs[:50]])  # First 50 paragraphs
            
            # Try to infer topic from keywords if not found in header
            if not metadata['topic']:
                metadata['topic'] = self._infer_topic_from_text(full_text)
            
            # Try to infer document type from content if not found in header
            if not metadata['document_type']:
                metadata['document_type'] = self._infer_doc_type_from_content(full_text)
            
            return metadata
        
        except Exception as e:
            logger.debug(f"Error extracting DOCX metadata from {Path(file_path).name}: {e}")
            return {'topic': None, 'document_type': None}
    
    def extract_metadata_from_pdf(self, file_path: str) -> Dict[str, Optional[str]]:
        """Extract Topic and Document Type from PDF metadata cover page (primary) or content (fallback)."""
        try:
            reader = PdfReader(file_path)
            metadata = {'topic': None, 'document_type': None}
            
            if len(reader.pages) == 0:
                return metadata
            
            # PRIMARY: Try to extract from first page (metadata cover page)
            first_page_text = reader.pages[0].extract_text()
            
            if first_page_text:
                lines = first_page_text.split('\n')
                
                for line in lines:
                    line = line.strip()
                    
                    # Look for "Topic:" followed by the topic value
                    if line.startswith('Topic:'):
                        topic = line.replace('Topic:', '').strip()
                        if topic and topic != 'None':
                            metadata['topic'] = topic
                    
                    # Look for "Document Type:" followed by the type
                    if line.startswith('Document Type:'):
                        doc_type = line.replace('Document Type:', '').strip()
                        if doc_type and doc_type != 'None':
                            metadata['document_type'] = doc_type
                
                # If we got both from cover page, return
                if metadata['topic'] and metadata['document_type']:
                    return metadata
            
            # FALLBACK: Analyze PDF content if header incomplete
            # Extract text from pages 2-5 (skip cover page if it exists, get main content)
            content_text = ""
            start_page = 1 if len(reader.pages) > 1 else 0
            for page_num in range(start_page, min(start_page + 5, len(reader.pages))):
                page_text = reader.pages[page_num].extract_text()
                if page_text:
                    content_text += page_text + "\n"
            
            # Try to infer topic from keywords if not found in header
            if not metadata['topic']:
                metadata['topic'] = self._infer_topic_from_text(content_text)
            
            # Try to infer document type from content if not found in header
            if not metadata['document_type']:
                metadata['document_type'] = self._infer_doc_type_from_content(content_text)
            
            return metadata
        
        except Exception as e:
            logger.debug(f"Error extracting PDF metadata from {Path(file_path).name}: {e}")
            return {'topic': None, 'document_type': None}
    
    def _infer_topic_from_text(self, text: str) -> Optional[str]:
        """Infer topic from text by searching for keyword matches."""
        text_lower = text.lower()
        
        # Score each topic based on keyword matches
        topic_scores = {}
        for topic, keywords in self.topic_keywords.items():
            score = sum(1 for keyword in keywords if keyword in text_lower)
            if score > 0:
                topic_scores[topic] = score
        
        # Return the topic with the highest score
        if topic_scores:
            return max(topic_scores.items(), key=lambda x: x[1])[0]
        
        return None
    
    def get_document_id(self, filename: str) -> Optional[str]:
        """Get document ID from filename by looking it up in the database."""
        conn = self.get_connection()
        try:
            result = conn.execute(
                "SELECT id FROM documents WHERE filename = ?",
                (filename,)
            ).fetchone()
            return result['id'] if result else None
        finally:
            conn.close()
    
    def update_document_topic(self, doc_id: str, topic: str) -> bool:
        """Update document_topics table with extracted topic."""
        if not topic:
            return False
        
        conn = self.get_connection()
        try:
            # Validate topic is in predefined list
            if topic not in PREDEFINED_TOPICS:
                logger.debug(f"Topic '{topic}' not in predefined topics, skipping")
                return False
            
            # Remove existing topics for this document
            conn.execute("DELETE FROM document_topics WHERE doc_id = ?", (doc_id,))
            
            # Insert the new topic
            conn.execute(
                "INSERT INTO document_topics (doc_id, topic) VALUES (?, ?)",
                (doc_id, topic)
            )
            conn.commit()
            return True
        except Exception as e:
            logger.error(f"Error updating topic for {doc_id}: {e}")
            return False
        finally:
            conn.close()
    
    def update_document_type(self, doc_id: str, doc_type: str) -> bool:
        """Update document_types table with extracted document type."""
        if not doc_type:
            return False
        
        conn = self.get_connection()
        try:
            # Validate doc_type is in predefined list
            if doc_type not in PREDEFINED_DOCUMENT_TYPES:
                logger.debug(f"Document type '{doc_type}' not in predefined types, skipping")
                return False
            
            # Remove existing type for this document
            conn.execute("DELETE FROM document_types WHERE doc_id = ?", (doc_id,))
            
            # Insert the new document type (column name is "doc_type")
            conn.execute(
                "INSERT INTO document_types (doc_id, doc_type) VALUES (?, ?)",
                (doc_id, doc_type)
            )
            conn.commit()
            return True
        except Exception as e:
            logger.error(f"Error updating document type for {doc_id}: {e}")
            return False
        finally:
            conn.close()
    
    def process_documents(self) -> Tuple[int, int, int, list]:
        """Process all documents and extract/update metadata."""
        if not os.path.exists(self.docs_folder):
            logger.error(f"Documents folder not found: {self.docs_folder}")
            return 0, 0, 0, []
        
        logger.info(f"Starting metadata extraction from documents in {self.docs_folder}")
        logger.info("Using primary extraction (headers) with fallback to content analysis")
        
        for filename in sorted(os.listdir(self.docs_folder)):
            file_path = os.path.join(self.docs_folder, filename)
            
            # Skip directories
            if os.path.isdir(file_path):
                continue
            
            # Get document ID from database
            doc_id = self.get_document_id(filename)
            if not doc_id:
                logger.debug(f"Document not found in database: {filename}")
                self.skipped_count += 1
                continue
            
            metadata = None
            
            # Extract metadata based on file type
            if filename.lower().endswith('.docx'):
                metadata = self.extract_metadata_from_docx(file_path)
            elif filename.lower().endswith('.pdf'):
                metadata = self.extract_metadata_from_pdf(file_path)
            else:
                self.skipped_count += 1
                continue
            
            # Update database with extracted metadata
            if metadata:
                try:
                    topic_updated = self.update_document_topic(doc_id, metadata.get('topic'))
                    type_updated = self.update_document_type(doc_id, metadata.get('document_type'))
                    
                    if topic_updated or type_updated:
                        topic_str = metadata.get('topic', 'N/A')
                        type_str = metadata.get('document_type', 'N/A')
                        logger.info(f"✓ {filename[:50]}")
                        logger.debug(f"  Topic: {topic_str}, Type: {type_str}")
                        self.processed_count += 1
                    else:
                        logger.debug(f"⊘ No valid metadata extracted: {filename}")
                        self.skipped_count += 1
                
                except Exception as e:
                    logger.error(f"✗ Error updating database for {filename}: {e}")
                    self.failed_count += 1
                    self.failed_files.append((filename, str(e)))
        
        return self.processed_count, self.skipped_count, self.failed_count, self.failed_files
    
    def print_summary(self):
        """Print processing summary."""
        print("\n" + "=" * 75)
        print("METADATA EXTRACTION FROM DOCUMENTS SUMMARY")
        print("=" * 75)
        print(f"Successfully processed: {self.processed_count} files")
        print(f"  (Extracted and updated database with Topic and Document Type)")
        print(f"Skipped: {self.skipped_count} files")
        print(f"Failed: {self.failed_count} files")
        
        if self.failed_files:
            print(f"\nFailed files:")
            for filename, error in self.failed_files:
                print(f"  - {filename}")
                print(f"    Error: {error}")
        
        print("\n" + "=" * 75)
        print("EXTRACTION METHOD:")
        print("  PRIMARY: Read Topic and Document Type from document headers")
        print("  FALLBACK: Analyze document content if headers missing/incomplete")
        print("    - Keyword matching for topic inference")
        print("    - Content pattern analysis for document type inference")
        print("=" * 75)
        print("\nThe document_topics and document_types tables have been updated.")
        print("The pill-style category filters and search will now use this data.")
        print("=" * 75)


def main():
    """Main entry point."""
    extractor = MetadataExtractor(str(DB_PATH), DOC_FOLDER)
    processed, skipped, failed, failed_files = extractor.process_documents()
    extractor.print_summary()
    
    return failed == 0


if __name__ == '__main__':
    import sys
    success = main()
    sys.exit(0 if success else 1)
