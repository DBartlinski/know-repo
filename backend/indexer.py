"""
Document indexer: scans DOC_FOLDER for supported files, extracts text and metadata,
and stores everything in a SQLite FTS5 virtual table for fast BM25 search.
"""

import hashlib
import logging
import sqlite3
from datetime import datetime, timezone
from pathlib import Path
from typing import Optional

from config import DB_PATH, DOC_FOLDER, SUPPORTED_EXTENSIONS, PREDEFINED_TOPICS, PREDEFINED_DOCUMENT_TYPES, TOPIC_MIN_COUNT

logger = logging.getLogger(__name__)


# ---------------------------------------------------------------------------
# Database setup
# ---------------------------------------------------------------------------

SCHEMA_SQL = """
CREATE TABLE IF NOT EXISTS documents (
    id          TEXT PRIMARY KEY,
    filename    TEXT NOT NULL,
    filepath    TEXT NOT NULL UNIQUE,
    filetype    TEXT NOT NULL,
    title       TEXT,
    author      TEXT,
    filesize    INTEGER,
    indexed_at  TEXT NOT NULL
);

CREATE VIRTUAL TABLE IF NOT EXISTS documents_fts USING fts5(
    doc_id UNINDEXED,
    title,
    content,
    tokenize='unicode61'
);

CREATE TABLE IF NOT EXISTS index_log (
    ran_at      TEXT NOT NULL,
    files_added INTEGER,
    files_updated INTEGER,
    files_removed INTEGER,
    duration_ms INTEGER
);

CREATE TABLE IF NOT EXISTS topics (
    topic       TEXT PRIMARY KEY,
    doc_count   INTEGER DEFAULT 0
);

CREATE TABLE IF NOT EXISTS document_topics (
    doc_id      TEXT NOT NULL,
    topic       TEXT NOT NULL,
    PRIMARY KEY (doc_id, topic),
    FOREIGN KEY (doc_id) REFERENCES documents(id) ON DELETE CASCADE,
    FOREIGN KEY (topic) REFERENCES topics(topic) ON DELETE CASCADE
);

CREATE TABLE IF NOT EXISTS document_types (
    doc_id      TEXT NOT NULL,
    doc_type    TEXT NOT NULL,
    PRIMARY KEY (doc_id, doc_type),
    FOREIGN KEY (doc_id) REFERENCES documents(id) ON DELETE CASCADE
);

CREATE TABLE IF NOT EXISTS document_corrections (
    doc_id          TEXT PRIMARY KEY,
    manual_type     TEXT,
    notes           TEXT,
    updated_at      TEXT,
    FOREIGN KEY (doc_id) REFERENCES documents(id) ON DELETE CASCADE
);
"""


def get_connection() -> sqlite3.Connection:
    conn = sqlite3.connect(str(DB_PATH))
    conn.row_factory = sqlite3.Row
    conn.execute("PRAGMA journal_mode=WAL")
    conn.execute("PRAGMA foreign_keys=ON")
    return conn


def init_db() -> None:
    with get_connection() as conn:
        conn.executescript(SCHEMA_SQL)


# ---------------------------------------------------------------------------
# Text extraction helpers
# ---------------------------------------------------------------------------

def _extract_pdf(path: Path) -> tuple[str, dict]:
    """Extract text and metadata from a PDF file."""
    metadata = {}
    try:
        import fitz  # PyMuPDF
        doc = fitz.open(str(path))
        meta = doc.metadata or {}
        metadata["title"] = meta.get("title") or ""
        metadata["author"] = meta.get("author") or ""
        pages = []
        for page in doc:
            pages.append(page.get_text("text"))
        doc.close()
        return "\n".join(pages), metadata
    except Exception as e:
        logger.warning("PyMuPDF failed for %s: %s — trying pdfminer", path.name, e)
        try:
            from pdfminer.high_level import extract_text as pdfminer_extract
            text = pdfminer_extract(str(path))
            return text or "", metadata
        except Exception as e2:
            logger.error("pdfminer also failed for %s: %s", path.name, e2)
            return "", metadata


def _extract_docx(path: Path) -> tuple[str, dict]:
    """Extract text and metadata from a DOCX file."""
    metadata = {}
    try:
        from docx import Document
        doc = Document(str(path))
        core = doc.core_properties
        metadata["title"] = core.title or ""
        metadata["author"] = core.author or ""
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        # Also grab text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        paragraphs.append(cell.text)
        return "\n".join(paragraphs), metadata
    except Exception as e:
        logger.error("python-docx failed for %s: %s", path.name, e)
        return "", metadata


def _extract_doc(path: Path) -> tuple[str, dict]:
    """Best-effort text extraction from legacy binary .doc files."""
    # python-docx cannot read .doc; attempt anyway (will raise, return empty)
    try:
        return _extract_docx(path)
    except Exception:
        logger.warning("Cannot extract text from legacy .doc: %s", path.name)
        return "", {}


def _extract_xlsx(path: Path) -> tuple[str, dict]:
    """Extract text from an XLSX file."""
    try:
        import openpyxl
        wb = openpyxl.load_workbook(str(path), read_only=True, data_only=True)
        rows = []
        for sheet in wb.worksheets:
            rows.append(f"[Sheet: {sheet.title}]")
            for row in sheet.iter_rows(values_only=True):
                cells = [str(c) for c in row if c is not None]
                if cells:
                    rows.append(" | ".join(cells))
        wb.close()
        return "\n".join(rows), {}
    except Exception as e:
        logger.error("openpyxl failed for %s: %s", path.name, e)
        return "", {}


def _extract_xls(path: Path) -> tuple[str, dict]:
    """Extract text from a legacy XLS file."""
    try:
        import xlrd
        wb = xlrd.open_workbook(str(path))
        rows = []
        for sheet in wb.sheets():
            rows.append(f"[Sheet: {sheet.name}]")
            for rx in range(sheet.nrows):
                cells = [str(sheet.cell_value(rx, cx)) for cx in range(sheet.ncols)]
                if any(c.strip() for c in cells):
                    rows.append(" | ".join(cells))
        return "\n".join(rows), {}
    except Exception as e:
        logger.error("xlrd failed for %s: %s", path.name, e)
        return "", {}


def extract_text(path: Path) -> tuple[str, dict]:
    ext = path.suffix.lower()
    if ext == ".pdf":
        return _extract_pdf(path)
    elif ext == ".docx":
        return _extract_docx(path)
    elif ext == ".doc":
        return _extract_doc(path)
    elif ext == ".xlsx":
        return _extract_xlsx(path)
    elif ext == ".xls":
        return _extract_xls(path)
    return "", {}


# ---------------------------------------------------------------------------
# Hashing / ID
# ---------------------------------------------------------------------------

def file_id(path: Path) -> str:
    """Stable document ID derived from the file path."""
    return hashlib.sha1(str(path.resolve()).encode()).hexdigest()


# ---------------------------------------------------------------------------
# Indexing
# ---------------------------------------------------------------------------

def _get_existing_ids(conn: sqlite3.Connection) -> dict[str, str]:
    """Return {doc_id: filepath} for all currently indexed documents."""
    rows = conn.execute("SELECT id, filepath FROM documents").fetchall()
    return {r["id"]: r["filepath"] for r in rows}


def _remove_document(conn: sqlite3.Connection, doc_id: str) -> None:
    conn.execute("DELETE FROM documents WHERE id = ?", (doc_id,))
    conn.execute("DELETE FROM documents_fts WHERE doc_id = ?", (doc_id,))
    conn.execute("DELETE FROM document_topics WHERE doc_id = ?", (doc_id,))
    conn.execute("DELETE FROM document_types WHERE doc_id = ?", (doc_id,))


def _extract_topics(title: str, filename: str) -> set[str]:
    """Extract topics from document title and filename using predefined topic list."""
    search_text = (title + " " + filename).lower()
    found_topics = set()
    
    # Match against predefined topics (case-insensitive, whole-phrase matching)
    for topic in PREDEFINED_TOPICS:
        topic_lower = topic.lower()
        # Check for exact phrase match or hyphen-separated variant matches
        if topic_lower in search_text:
            found_topics.add(topic)
        # Also check without special characters for flexible matching
        elif topic_lower.replace(" - ", " ").replace(" – ", " ") in search_text.replace(" - ", " ").replace(" – ", " "):
            found_topics.add(topic)
    
    return found_topics


def _extract_document_type(title: str, filename: str) -> str:
    """Extract document type from filename patterns and context."""
    search_text = (title + " " + filename).upper()
    
    # Pattern matching rules (ordered by specificity)
    if "BRIEFING PAPER" in search_text:
        if "CONGRESS" in search_text:
            return "Briefing Papers - Congress"
        elif "VA" in search_text or "VHA" in search_text:
            return "Briefing Papers - VA-VHA"
        else:
            return "Briefing Papers - Other"
    
    if "TALKING POINTS" in search_text or "TALKING POINT" in search_text:
        if "CONGRESS" in search_text:
            return "Talking Points - Congress"
        elif "VA" in search_text or "VHA" in search_text:
            return "Talking Points - VA-VHA"
        else:
            return "Talking Points - Other"
    
    if "TESTIMONY" in search_text or "CONGRESS SPEECH" in search_text:
        return "Testimony - Congress Speeches"
    
    if "RFI--" in filename or "RFI_" in filename or ("RFI" in search_text and "CONGRESS" in search_text):
        return "RFI - Congress"
    
    if "MANDATED" in search_text or "CONGRESSIONAL REPORT" in search_text:
        return "Mandated Congressional Reports"
    
    if "NEWS" in search_text or "ARTICLE" in search_text:
        return "News Articles"
    
    if "PRESS RELEASE" in search_text:
        return "Press Releases"
    
    if "SLIDE" in search_text or "PRESENTATION" in search_text or filename.endswith(".pptx"):
        return "Slide Presentations"
    
    if "PUBLICATION" in search_text or "STUDY" in search_text:
        return "Publications-Studies"
    
    if "RFA" in search_text and "RFI" not in search_text:
        return "RFA"
    
    if "GAO" in search_text:
        return "GAO"
    
    if "BUDGET" in search_text:
        return "Budget Justification"
    
    # Fallback: Other
    return "Other-Misc Documents"


def _upsert_document(conn: sqlite3.Connection, path: Path) -> str:
    """Index or re-index a single file. Returns 'added' or 'updated'."""
    doc_id = file_id(path)
    text, meta = extract_text(path)
    title = meta.get("title") or path.stem
    author = meta.get("author") or ""
    now = datetime.now(timezone.utc).isoformat()

    existing = conn.execute(
        "SELECT id FROM documents WHERE id = ?", (doc_id,)
    ).fetchone()

    if existing:
        conn.execute(
            "UPDATE documents SET filename=?, filepath=?, filetype=?, title=?, author=?, filesize=?, indexed_at=? WHERE id=?",
            (path.name, str(path.resolve()), path.suffix.lower().lstrip("."),
             title, author, path.stat().st_size, now, doc_id),
        )
        conn.execute("DELETE FROM documents_fts WHERE doc_id = ?", (doc_id,))
        conn.execute("DELETE FROM document_topics WHERE doc_id = ?", (doc_id,))
        conn.execute("DELETE FROM document_types WHERE doc_id = ?", (doc_id,))
        status = "updated"
    else:
        conn.execute(
            "INSERT INTO documents (id, filename, filepath, filetype, title, author, filesize, indexed_at) VALUES (?,?,?,?,?,?,?,?)",
            (doc_id, path.name, str(path.resolve()), path.suffix.lower().lstrip("."),
             title, author, path.stat().st_size, now),
        )
        status = "added"

    conn.execute(
        "INSERT INTO documents_fts (doc_id, title, content) VALUES (?, ?, ?)",
        (doc_id, title, text),
    )
    
    # Extract and store topics
    topics = _extract_topics(title, path.name)
    for topic in topics:
        conn.execute("INSERT OR IGNORE INTO topics (topic) VALUES (?)", (topic,))
        conn.execute(
            "INSERT OR IGNORE INTO document_topics (doc_id, topic) VALUES (?, ?)",
            (doc_id, topic),
        )
    
    # Extract and store document type
    doc_type = _extract_document_type(title, path.name)
    conn.execute(
        "INSERT OR IGNORE INTO document_types (doc_id, doc_type) VALUES (?, ?)",
        (doc_id, doc_type),
    )
    
    return status


def run_index(doc_folder: Optional[Path] = None) -> dict:
    """
    Scan the document folder and update the SQLite index.
    Returns a summary dict with counts.
    """
    folder = doc_folder or DOC_FOLDER
    if not folder.exists():
        raise FileNotFoundError(f"DOC_FOLDER does not exist: {folder}")

    init_db()
    start = datetime.now(timezone.utc)
    added = updated = removed = 0

    with get_connection() as conn:
        existing = _get_existing_ids(conn)
        found_ids = set()

        for path in folder.rglob("*"):
            if path.suffix.lower() not in SUPPORTED_EXTENSIONS:
                continue
            if not path.is_file():
                continue

            doc_id = file_id(path)
            found_ids.add(doc_id)
            try:
                status = _upsert_document(conn, path)
                if status == "added":
                    added += 1
                else:
                    updated += 1
            except Exception as e:
                logger.error("Failed to index %s: %s", path.name, e)

        # Remove documents whose files no longer exist
        for old_id in set(existing) - found_ids:
            _remove_document(conn, old_id)
            removed += 1

        duration_ms = int(
            (datetime.now(timezone.utc) - start).total_seconds() * 1000
        )
        conn.execute(
            "INSERT INTO index_log (ran_at, files_added, files_updated, files_removed, duration_ms) VALUES (?,?,?,?,?)",
            (start.isoformat(), added, updated, removed, duration_ms),
        )

    logger.info("Index complete: +%d updated=%d removed=%d (%dms)", added, updated, removed, duration_ms)
    return {
        "files_added": added,
        "files_updated": updated,
        "files_removed": removed,
        "duration_ms": duration_ms,
        "ran_at": start.isoformat(),
    }
